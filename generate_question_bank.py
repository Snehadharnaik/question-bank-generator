# generate_question_bank.py
import re
import io
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import pdfplumber

from sklearn.feature_extraction.text import TfidfVectorizer

# ------------------ Utilities ------------------

def clean_text_for_matching(text: str) -> str:
    if not text:
        return ""
    t = re.sub(r'\s+', ' ', text)  # collapse whitespace
    t = t.strip().lower()
    return t

def get_words(text):
    return re.findall(r"\b[a-zA-Z]+\b", text.lower())

# ------------------ Syllabus parsing & technical term extraction ------------------

def extract_text_from_docx_bytes(b):
    # b is BytesIO or file-like
    doc = Document(b)
    texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(texts)

def extract_text_from_pdf_bytes(b):
    text = ""
    with pdfplumber.open(b) as pdf:
        for page in pdf.pages:
            ptext = page.extract_text()
            if ptext:
                text += ptext + "\n"
    return text

def build_syllabus_terms(syllabus_file):
    """
    Build a set of candidate technical terms (1-3 word ngrams) from syllabus text.
    Returns a set of cleaned phrases (lowercase).
    """
    if syllabus_file is None:
        return set()

    name = syllabus_file.name.lower()
    raw_text = ""
    try:
        if name.endswith(".docx"):
            raw_text = extract_text_from_docx_bytes(syllabus_file)
        elif name.endswith(".pdf"):
            raw_text = extract_text_from_pdf_bytes(syllabus_file)
        else:
            # unsupported, return empty
            return set()
    except Exception:
        # fallback: try reading raw bytes and decode
        try:
            raw_text = str(syllabus_file.read(), errors="ignore")
        except Exception:
            raw_text = ""

    raw_text = clean_text_for_matching(raw_text)
    words = get_words(raw_text)

    # Filter candidate words: length >=4 (to avoid short common words)
    filt_words = [w for w in words if len(w) >= 3]

    # Build ngrams (1..3) contiguous from syllabus text and add to set if not stopwords
    # We'll also filter out very common words by using an english stopword list (scikit-learn's)
    from sklearn.feature_extraction import text as sktext
    stopset = set(sktext.ENGLISH_STOP_WORDS)

    terms = set()
    n = len(filt_words)
    for i in range(n):
        for L in (3, 2, 1):  # prefer longer phrases first
            if i + L <= n:
                phrase_words = filt_words[i:i+L]
                phrase = " ".join(phrase_words)
                # filter: must contain at least one non-stopword and length>3 overall
                if any(w not in stopset for w in phrase_words):
                    # discard phrase if mostly numeric or punctuation (already removed)
                    terms.add(clean_text_for_matching(phrase))
    # deduplicate and return
    return set([t for t in terms if len(t) >= 2])

# ------------------ Keyword extraction ------------------

def find_syllabus_keyword_in_question(question_text, syllabus_terms):
    """
    Return the longest matching syllabus term present in the question (case-insensitive).
    We check phrases in descending length order (so multi-word phrases preferred).
    """
    q = clean_text_for_matching(question_text)
    if not q or not syllabus_terms:
        return None

    # Sort syllabus terms by word count and length so longer multi-word matches checked first
    sorted_terms = sorted(syllabus_terms, key=lambda x: (-len(x.split()), -len(x)))
    for term in sorted_terms:
        # match whole phrase (simple substring check is usually good for these texts)
        if term in q:
            return term
    return None

def tfidf_top_ngram_for_question(all_questions, q_index, ngram_range=(1,3), top_k=1):
    """
    Fit TF-IDF on all_questions (list[str]) and return top ngrams for question at q_index.
    Returns a string (top ngram) or None.
    """
    # Vectorizer: words, 1..3 grams, English stop words removed
    vectorizer = TfidfVectorizer(ngram_range=ngram_range, stop_words='english', token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z]+\b', max_features=5000)
    try:
        X = vectorizer.fit_transform(all_questions)
    except ValueError:
        return None
    feature_names = vectorizer.get_feature_names_out()
    row = X[q_index]
    if row.nnz == 0:
        return None
    # Get top indices by tfidf score
    import numpy as np
    coo = row.tocoo()
    tuples = list(zip(coo.col, coo.data))
    if not tuples:
        return None
    tuples.sort(key=lambda x: -x[1])
    top_feats = [feature_names[idx] for idx, _ in tuples[:top_k]]
    return top_feats[0] if top_feats else None

def extract_keyword_for_question(question_text, idx, all_questions, syllabus_terms):
    """
    1) Try to find multi-word technical phrase from syllabus_terms (prefer 3-word > 2-word >1-word)
    2) If not found, use TF-IDF fallback across all_questions to pick top 1-3gram.
    3) Post-process: return keyword as cleaned text.
    """
    # 1. syllabus-based
    kw = find_syllabus_keyword_in_question(question_text, syllabus_terms)
    if kw:
        return kw

    # 2. TF-IDF fallback
    tt = tfidf_top_ngram_for_question(all_questions, idx, ngram_range=(1,3), top_k=1)
    if tt:
        return tt

    # 3. final fallback: first meaningful long word
    words = re.findall(r'\b[a-zA-Z]{4,}\b', question_text.lower())
    if words:
        return words[0]
    return "General"

# ------------------ Helpers (bloom, difficulty, qtype) ------------------

def detect_bloom_level(question):
    question = (question or "").lower()
    bloom_keywords = {
        "L1": ["define", "list", "name", "state"],
        "L2": ["explain", "describe", "summarize", "classify"],
        "L3": ["solve", "use", "demonstrate", "compute", "calculate", "determine"],
        "L4": ["compare", "differentiate", "analyze", "distinguish"],
        "L5": ["justify", "evaluate", "assess", "argue"],
        "L6": ["design", "develop", "formulate", "construct"]
    }
    for level, verbs in bloom_keywords.items():
        for verb in verbs:
            if verb in question:
                return level
    return "L2"

def assign_difficulty(bloom_level):
    return {
        "L1": "Low",
        "L2": "Low",
        "L3": "Medium",
        "L4": "Medium",
        "L5": "High",
        "L6": "High"
    }.get(bloom_level, "Medium")

def classify_question_type(question):
    q = (question or "").lower()
    return "P" if any(word in q for word in ["calculate", "solve", "determine", "find"]) else "T"

# ------------------ DOCX generation ------------------

def add_table_rows_centered(doc, rows):
    """
    rows: list of (label, value)
    Adds a 2-column table centered in the document, sets font and returns table object.
    """
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for label, value in rows:
        r_cells = table.add_row().cells
        r_cells[0].text = str(label)
        r_cells[1].text = str(value)
        for cell in r_cells:
            for p in cell.paragraphs:
                # ensure run exists
                if len(p.runs) == 0:
                    p.add_run()
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    return table

def generate_docx_from_dataframe(df, syllabus_unit_map, syllabus_terms, output_path="QuestionBank_Output.docx"):
    """
    df: pandas DataFrame expected columns at minimum: Question, Unit, Subunit, Marks, Answer, Teacher ID, Tag, CO
    syllabus_unit_map: dict mapping unit number -> unit name (optional)
    syllabus_terms: set of phrases extracted from syllabus
    """
    df = df.fillna("")
    doc = Document()

    # Prepare questions list for TF-IDF
    all_questions = [str(q) for q in df["Question"].astype(str).tolist()]

    for idx, row in df.reset_index(drop=True).iterrows():
        qno = idx + 1
        question = str(row.get("Question", "")).strip()
        unit = str(row.get("Unit", "")).strip()
        subunit = str(row.get("Subunit", "")).strip()
        marks = str(row.get("Marks", "")).strip()
        answer = str(row.get("Answer", "")).strip()
        teacher_id = str(row.get("Teacher ID", "")).strip()
        tag = str(row.get("Tag", "")).strip()
        co = str(row.get("CO", "")).strip()

        bloom = detect_bloom_level(question)
        difficulty = assign_difficulty(bloom)
        qtype = classify_question_type(question)

        keyword = extract_keyword_for_question(question, idx, all_questions, syllabus_terms)

        unit_name = syllabus_unit_map.get(unit, "") if unit else ""
        display_tag = tag  # per your request, Tag from CSV is used as-is

        rows = [
            ("Question No.", qno),
            ("Question", question),
            ("Unit", f"Unit {unit}" if unit else ""),
            ("Subunit", subunit),
            ("Marks", marks),
            ("Difficulty", difficulty[0].upper() if difficulty else ""),
            ("Answer", answer),
            ("Question Type", qtype),
            ("Tag", display_tag),
            ("Keywords", keyword),
            ("Blooms Taxonomy", bloom),
            ("Course Outcome", co if co else "CO1"),
            ("Teacher ID", f"<{teacher_id}>"),
            ("Year", "<System updates>"),
            ("Year asked", "<System updates>"),
            ("Frequency", "<System updates>")
        ]

        add_table_rows_centered(doc, rows)
        doc.add_page_break()

    doc.save(output_path)
    return output_path

# ------------------ Syllabus unit mapping ------------------

def read_unit_mapping_from_docx_bytes(b):
    """
    Parse docx bytes for lines starting with '1 <UnitName>' or similar to map unit numbers to names.
    Returns dict unit_num_str -> unit name
    """
    try:
        doc = Document(b)
    except Exception:
        return {}
    mapping = {}
    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue
        # common pattern: "1\tWATER SUPPLY SYSTEM" or "1  WATER SUPPLY SYSTEM"
        m = re.match(r'^(\d+)[\s\.\-\)]*\s*(.+)$', line)
        if m:
            unitnum = m.group(1).strip()
            name = m.group(2).strip()
            # heuristic: skip very short names
            if len(name) > 2:
                mapping[unitnum] = name
    return mapping

def read_unit_mapping_from_pdf_bytes(b):
    mapping = {}
    try:
        with pdfplumber.open(b) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    m = re.match(r'^(\d+)[\s\.\-\)]*\s*(.+)$', line)
                    if m:
                        unitnum = m.group(1).strip()
                        name = m.group(2).strip()
                        if len(name) > 2:
                            mapping[unitnum] = name
    except Exception:
        pass
    return mapping

# ------------------ Streamlit UI ------------------

st.set_page_config(page_title="Question Bank Generator", layout="centered")
st.title("📚 Question Bank Generator")

st.markdown("Upload CSV (with columns: `Question,Unit,Subunit,Marks,Answer,Teacher ID,Tag,CO`) and optional syllabus (DOCX/PDF).")

csv_file = st.file_uploader("Upload Questions CSV", type=["csv"])
syllabus_file = st.file_uploader("Upload Syllabus (optional) - .docx or .pdf", type=["docx","pdf"])

syllabus_terms = set()
unit_map = {}

if syllabus_file:
    # syllabus_file is a UploadedFile object (has .read())
    try:
        # For docx
        if syllabus_file.name.lower().endswith(".docx"):
            # need a BytesIO for multiple reads
            b = io.BytesIO(syllabus_file.getvalue())
            raw_text = extract_text_from_docx_bytes(b)
            # build terms
            syllabus_terms = build_syllabus_terms(b)
            # and unit mapping
            b2 = io.BytesIO(syllabus_file.getvalue())
            unit_map = read_unit_mapping_from_docx_bytes(b2)
        elif syllabus_file.name.lower().endswith(".pdf"):
            b = io.BytesIO(syllabus_file.getvalue())
            raw_text = extract_text_from_pdf_bytes(b)
            syllabus_terms = build_syllabus_terms(b)
            b2 = io.BytesIO(syllabus_file.getvalue())
            unit_map = read_unit_mapping_from_pdf_bytes(b2)
    except Exception as e:
        st.warning(f"Could not fully parse syllabus file ({e}). Syllabus-based keywords may be limited.")
        syllabus_terms = set()
        unit_map = {}

if csv_file:
    try:
        df = pd.read_csv(csv_file)
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        df = None

    if df is not None:
        # Ensure required columns exist or create defaults
        for c in ["Question","Unit","Subunit","Marks","Answer","Teacher ID","Tag","CO"]:
            if c not in df.columns:
                df[c] = ""

        st.write("Preview (first 5 rows):")
        st.dataframe(df.head())

        if st.button("Generate Question Bank (.docx)"):
            out_path = generate_docx_from_dataframe(df, unit_map, syllabus_terms, output_path="QuestionBank_Output.docx")
            with open(out_path, "rb") as f:
                st.download_button("Download generated DOCX", f, file_name="QuestionBank_Output.docx")
            st.success("Generated: QuestionBank_Output.docx")
else:
    st.info("Please upload the questions CSV to begin.")
